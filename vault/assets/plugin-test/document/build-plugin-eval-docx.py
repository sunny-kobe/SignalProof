from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "outputs"
OUT.mkdir(parents=True, exist_ok=True)


def set_run(run, *, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(8)
run = title.add_run("Codex 推荐插件真实试跑评估")
set_run(run, size=22, bold=True, color="0B2545")

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
run = subtitle.add_run("SignalProof 内部流程资产化摘要 | 2026-06-21")
set_run(run, size=11, color="555555")

summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(10)
run = summary.add_run(
    "本文件用于试跑 Documents 插件能力：从插件审计结果生成可交付 DOCX，"
    "并通过结构化检查确认内容、标题层级和表格数据存在。"
)
set_run(run)

doc.add_heading("结论摘要", level=1)
items = [
    "强增益：GitHub、Record & Replay、Spreadsheets、Data Visualization。",
    "条件增益：Hugging Face、last30days、Zotero、Sentry/PostHog/Mixpanel/Amplitude/Datadog。",
    "暂不默认：Browser、Computer Use 当前环境失败；app-only connector 缺少工具入口或账号授权。",
    "产物建议：表格和可视化默认进入工具账本；DOCX/PDF/PPTX 在正式交付件时启用。",
]
for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(item))

doc.add_heading("插件试跑矩阵", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
headers = ["插件组", "结果质量", "实际证据", "默认建议"]
widths = [1800, 1400, 3200, 2960]
for cell, header, width in zip(table.rows[0].cells, headers, widths):
    set_cell_shading(cell, "E8EEF5")
    set_cell_width(cell, width)
    set_run(cell.paragraphs[0].add_run(header), bold=True, color="0B2545")

rows = [
    ["GitHub / last30days / Hugging Face", "中到强", "公开仓库、最近 30 天讨论、模型/数据集 API", "调研阶段默认候选"],
    ["Record & Replay", "强", "event_stream_status 可读", "资产化阶段默认候选"],
    ["Spreadsheets / Data Visualization", "强", "XLSX、PNG、HTML 已生成", "工具账本与报告可视化默认候选"],
    ["Documents / PDF / Presentations", "中", "本轮生成 DOCX/PDF/PPTX；渲染依赖部分缺失", "正式交付件启用"],
    ["Readwise / Scite / Semrush 等 app-only", "阻塞", "manifest 为 app-only；当前线程未暴露 connector", "授权后专项跑，不默认"],
]
for row in rows:
    cells = table.add_row().cells
    for cell, value, width in zip(cells, row, widths):
        set_cell_width(cell, width)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(value), size=10)

doc.add_heading("边界", level=1)
p = doc.add_paragraph()
set_run(
    p.add_run(
        "本轮只能证明插件对个人内部工作流的增益和阻塞点，不能证明 SignalProof 已获得市场验证。"
    )
)

doc.save(OUT / "plugin-eval-brief.docx")
